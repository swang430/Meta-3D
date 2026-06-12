from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/site-debug/2026-06-12-fs16-onsite-work-plan.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 30, 45)
MUTED = RGBColor(91, 103, 119)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
CAUTION = "FFF2CC"
OK_FILL = "EAF4EA"

USABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(run, *, name="Calibri", east_asia="Microsoft YaHei", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_style_font(style, *, name="Calibri", east_asia="Microsoft YaHei", size=11, color=None, bold=None):
    font = style.font
    font.name = name
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    font.size = Pt(size)
    if color is not None:
        font.color.rgb = color
    if bold is not None:
        font.bold = bold


def set_paragraph_spacing(paragraph, *, before=0, after=6, line=1.25):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def paragraph_border_bottom(paragraph, color="9BB7D4", size="8", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, margins=CELL_MARGIN_DXA):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in margins.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(1, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[min(idx, len(widths_dxa) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def cell_text(cell, text, *, bold=False, color=INK, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0, line=1.15)
    run = p.add_run(str(text))
    set_run_font(run, size=size, color=color, bold=bold)


def add_table(doc, headers, rows, widths_dxa, *, header_fill=LIGHT_BLUE, font_size=9.4):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        shade_cell(hdr.cells[idx], header_fill)
        cell_text(hdr.cells[idx], text, bold=True, color=DARK_BLUE, size=font_size)
    for row_values in rows:
        row = table.add_row()
        for idx, text in enumerate(row_values):
            cell_text(row.cells[idx], text, size=font_size)
    set_table_geometry(table, widths_dxa)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line=1.0)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_para(doc, text="", *, bold_label=None, after=6, color=INK):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=after, line=1.25)
    if bold_label:
        r = p.add_run(f"{bold_label}: ")
        set_run_font(r, size=11, color=color, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=11, color=color)
    return p


def add_callout(doc, title, body, *, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, {"top": 120, "bottom": 120, "start": 160, "end": 160})
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=0, after=2, line=1.2)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    set_paragraph_spacing(p2, before=0, after=0, line=1.2)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.2, color=INK)
    set_table_geometry(table, [USABLE_WIDTH_DXA - TABLE_INDENT_DXA])
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=6, line=1.0)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, size=11, color=INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(header, before=0, after=0, line=1.0)
    r = header.add_run("FS16 onsite debug guide")
    set_run_font(r, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(footer, before=0, after=0, line=1.0)
    r = footer.add_run("Meta-3D | 2026-06-12 | Internal onsite reference")
    set_run_font(r, size=9, color=MUTED)


def build_doc():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    set_paragraph_spacing(title, before=0, after=4, line=1.1)
    r = title.add_run("FS16 现场调试工作计划与操作步骤")
    set_run_font(r, size=22, color=INK, bold=True)

    sub = doc.add_paragraph()
    set_paragraph_spacing(sub, before=0, after=10, line=1.2)
    r = sub.add_run("目标: 真实 FS16 信道仿真器 + mock 基站 + mock DUT/KPI 的软件闭环确认")
    set_run_font(r, size=12.5, color=MUTED)

    meta_rows = [
        ("日期", "2026-06-12"),
        ("现场主目标", "通过 GUI/HAL 控制真实 PROPSIM FS16, 加载并运行 FS16 内部 .smu 文件。"),
        ("今日基站/DUT策略", "baseStation 使用 mock; DUT/终端 KPI 来自 mock BS/DUT。"),
        ("FS16主 endpoint", "TCPIP0::192.168.0.100::hislip0::INSTR"),
        ("默认 .smu", "Emulation0609.smu; 可在 GUI/Sequence Runner 中改为现场实际文件名或完整路径。"),
    ]
    add_table(doc, ["字段", "内容"], meta_rows, [1800, 7560], font_size=9.6)

    add_callout(
        doc,
        "出发前关键确认",
        "real/mock 模式由 GUI 的仪器 driver mode 和全局 HAL 模式决定; hybrid 序列只校验当前 HAL 已加载的 driver 是否符合参数意图, 不在后端强行创建 mock 基站。FS16 .smu 文件名/路径由 GUI connection_params 或 Sequence Runner 参数传入; 后端的 Emulation0609.smu 只是默认预填/兜底值。",
        fill=OK_FILL,
    )

    add_heading(doc, "1. 软件端检查结论", 1)
    add_table(
        doc,
        ["检查点", "结论", "证据/位置"],
        [
            (
                "real/mock 切换",
                "通过 GUI driver mode 写入 auto/mock/real; HAL 初始化时读取 DB 后决定使用 real driver 或 mock driver。",
                "gui/src/App.tsx 驱动模式 SegmentedControl; api-service/app/api/instrument.py driver-mode; instrument_hal_service.py _decide_use_real。",
            ),
            (
                "Hybrid 序列的 BS 模式",
                "base_station_mode 是 Sequence Runner 参数, 默认 mock, 可切 real。序列不创建/硬改 baseStation driver; 只在 mock 参数与 real BS driver 不一致时拦截。",
                "fs16_hybrid_kpi_smoke.py params_schema 与 run() 中 actual_bs_mode 校验。",
            ),
            (
                "FS16 文件名",
                "remote_playback_file 是 GUI/Sequence Runner 可编辑字段。Emulation0609.smu 是默认值, 不是不可改常量。",
                "gui/src/App.tsx FS16 Playback 文件输入; fs16_hybrid_kpi_smoke.py / fs16_playback_smoke.py params_schema。",
            ),
            (
                "FS16 路径",
                "可以填完整 D:\\... 路径; 只填文件名时 driver 按 playback_dir 组合路径。playback_dir 也可由 connection_params 覆盖。",
                "propsim_fs16_playback.py _configured_remote_file() 与 _remote_path()。",
            ),
            (
                "出发前测试",
                "后端模式/FS16 hybrid/playback driver 相关测试通过; 前端 build 通过。",
                "pytest: 32 passed; gui npm run build passed; Vite 仅有既有 chunk warning。",
            ),
        ],
        [2100, 3000, 4260],
        font_size=8.8,
    )

    add_heading(doc, "2. 现场总原则", 1)
    add_table(
        doc,
        ["原则", "执行方式"],
        [
            ("先软件闭环, 再真实传导", "先确认 GUI -> 后端 -> HAL -> FS16 load/start/stop -> mock KPI 展示。"),
            ("只让 FS16 走 real", "channelEmulator 选择 PROPSIM FS16 且 driver mode=Real; baseStation 明确 Mock 或 Auto+全局 Mock。"),
            ("避免误碰真实基站", "当前没有真实基站参与时, base_station_mode 保持 mock; 如果 HAL 已加载 real BS, hybrid 序列会拒绝执行。"),
            ("路径以现场文件为准", "优先使用 FS16 内已有 .smu; 现场确认文件名后在 GUI/Sequence Runner 输入。"),
            ("raw 5025 只诊断", "主路径使用 HiSLIP; raw 5025 仅在 HiSLIP 正常但需要排查 raw socket 时使用。"),
        ],
        [2300, 7060],
        font_size=9.2,
    )

    add_heading(doc, "3. GUI 配置确认", 1)
    add_heading(doc, "3.1 仪器资源配置", 2)
    add_table(
        doc,
        ["类别", "GUI 设置", "现场目标值", "备注"],
        [
            ("channelEmulator", "型号", "PROPSIM FS16", "绑定 RealPropsimFs16PlaybackDriver。"),
            ("channelEmulator", "驱动模式", "Real", "即使全局 HAL 是 Mock, per-instrument Real 仍会连接真实 FS16。不要使用 MOCK_FORCE。"),
            ("channelEmulator", "控制端点", "TCPIP0::192.168.0.100::hislip0::INSTR", "HiSLIP 主路径。"),
            ("baseStation", "驱动模式", "Mock", "今天不接真实基站仿真器。也可 Auto + 全局 Mock。"),
            ("DUT/终端", "指标来源", "Mock KPI", "结果只验证软件流程, 不作为真实射频性能结论。"),
        ],
        [1800, 1700, 3000, 2860],
        font_size=8.8,
    )

    add_heading(doc, "3.2 FS16 文件参数", 2)
    add_table(
        doc,
        ["字段", "在哪里改", "建议值", "说明"],
        [
            ("remote_playback_file", "仪器资源 FS16 Playback 文件; 或 Sequence Runner 参数", "Emulation0609.smu", "可改成现场实际 .smu 文件名或完整 D:\\ 路径。"),
            ("playback_dir", "connection_params JSON", "D:\\User Playbacks", "只填文件名时会拼到这个目录下; 如 FS16 目录不同可改。"),
            ("verify_remote_file_exists", "GUI switch 或 Sequence Runner 参数", "true", "load 前先用 FS16 目录查询确认文件可见。"),
            ("start_playback", "Sequence Runner 参数", "true", "load 成功后是否发送 start。"),
            ("stop_after_s / cleanup_on_finish", "Sequence Runner 参数", "5 / true", "默认现场 smoke 后自动停止, 避免 FS16 留在 running。"),
        ],
        [1900, 2700, 2100, 2660],
        font_size=8.8,
    )

    add_heading(doc, "4. 现场操作步骤", 1)
    add_heading(doc, "Phase 0: 开机与网络前置确认", 2)
    add_table(
        doc,
        ["顺序", "操作", "通过标准", "失败处理"],
        [
            ("0.1", "确认电脑网口/IP 与 FS16 同网段, FS16 地址为 192.168.0.100。", "能 ping 或至少 TCP/HiSLIP 可达。", "检查网线、IP、交换机、FS16 网口。"),
            ("0.2", "确认 FS16 UI 中目标 .smu 文件存在。", "能在 FS16 本机看到 Emulation0609.smu 或现场目标文件。", "把文件放入 FS16 指定目录, 记录真实文件名。"),
            ("0.3", "启动后端和前端, 打开仪器资源页面。", "GUI 可加载 catalog 和 lab profile。", "查看后端日志和浏览器控制台。"),
        ],
        [900, 3900, 2500, 2060],
        font_size=8.7,
    )

    add_heading(doc, "Phase 1: 配置仪器模式并 reload HAL", 2)
    add_table(
        doc,
        ["顺序", "操作", "通过标准", "记录"],
        [
            ("1.1", "channelEmulator 选择 PROPSIM FS16, endpoint 填 HiSLIP。", "控制端点显示 TCPIP0::192.168.0.100::hislip0::INSTR。", "记录 endpoint。"),
            ("1.2", "channelEmulator driver mode 选 Real。", "GUI 显示 Real; 保存后 reload HAL。", "记录 reload 时间。"),
            ("1.3", "baseStation driver mode 选 Mock。", "GUI 显示 Mock; 无真实 BS 连接动作。", "截图/记录。"),
            ("1.4", "HAL reload。", "系统就绪/driver 列表中 channelEmulator 是 FS16 real, baseStation 是 mock。", "若 FS16 未加载, 先做 SCPI probe。"),
        ],
        [900, 3900, 2500, 2060],
        font_size=8.7,
    )

    add_heading(doc, "Phase 2: FS16 基础连通与目录确认", 2)
    add_table(
        doc,
        ["顺序", "操作", "期望结果", "失败分支"],
        [
            ("2.1", "运行 FS16 health 或 SCPI probe: *IDN?。", "返回 Keysight/F8820A/FS16 相关身份信息。", "优先检查 HiSLIP endpoint; raw 5025 不作为主路径阻塞项。"),
            ("2.2", "查询 MMEM:CDIR?。", "目录为 D:\\User Playbacks 或现场实际 playback 目录。", "若目录不同, 更新 playback_dir 或直接填完整路径。"),
            ("2.3", "查询 MMEM:CAT? 或运行 hybrid 的 verify step。", "结果中包含目标 .smu 文件。", "文件不可见时先修正 FS16 文件位置/文件名。"),
        ],
        [900, 3600, 2700, 2160],
        font_size=8.7,
    )

    add_heading(doc, "Phase 3: FS16 playback smoke", 2)
    add_table(
        doc,
        ["顺序", "操作", "参数", "通过标准"],
        [
            ("3.1", "先跑 fs16_playback_smoke load-only。", "remote_playback_file=现场文件; start_playback=false; cleanup_on_finish=false。", "connect channelEmulator + FS16 load playback 成功。"),
            ("3.2", "再跑 fs16_playback_smoke start/stop。", "start_playback=true; stop_after_s=5; cleanup_on_finish=true。", "load/start/wait/stop 全部成功。"),
            ("3.3", "若 load 报 -300 或 corrupt/missing。", "记录 SYST:ERR? 与 DIAG:SIMU:STATe?。", "去 FS16 本机手动打开 .smu, 判断依赖文件/版本兼容。"),
        ],
        [850, 3100, 2800, 2610],
        font_size=8.6,
    )

    add_heading(doc, "Phase 4: Hybrid KPI smoke", 2)
    add_table(
        doc,
        ["参数", "现场建议值", "说明"],
        [
            ("remote_playback_file", "Emulation0609.smu 或现场确认文件", "可填完整路径。"),
            ("base_station_mode", "mock", "今天不接真实 BS; 必须与 HAL 中 baseStation mock 匹配。"),
            ("frequency_mhz", "3500 或现场虚拟配置", "mock BS 按此生成 attach/KPI 窗口。"),
            ("bandwidth_mhz / scs_khz / band", "100 / 30 / n78", "按本次虚拟基站参数填写。"),
            ("mimo_layers / dl_power_dbm", "2 / -50", "仅影响 mock KPI 模拟值。"),
            ("throughput_windows / window_s", "3 / 0.2", "现场 smoke 可保持默认。"),
        ],
        [2500, 2500, 4360],
        font_size=8.8,
    )
    add_table(
        doc,
        ["通过标准", "必须看到"],
        [
            ("FS16 链路", "connect channelEmulator (FS16); FS16 verify playback file; FS16 load playback; FS16 start playback; FS16 stop playback。"),
            ("Mock 基站链路", "connect baseStation (mock); BS set_cell_config; BS start_signaling; DUT attach。"),
            ("KPI 展示", "DL/UL throughput, DL/UL BLER, CQI, Rank Indicator, MCS, RSRP/SINR 如 mock 数据可用。"),
            ("来源标识", "结果卡显示 CE real / BS mock / DUT mock / KPI mock baseStation 或等价来源。"),
            ("清理状态", "默认应显示 FS16 playback stopped; BS signaling stopped。"),
        ],
        [2300, 7060],
        font_size=8.8,
    )

    add_heading(doc, "5. 不通过时的停止标准", 1)
    add_table(
        doc,
        ["现象", "立即停止/不继续的条件", "下一步"],
        [
            ("HAL 加载 real BS", "base_station_mode=mock 但实际 baseStation 是 real。", "不要继续; GUI 把 baseStation driver mode 改 Mock 并 reload HAL。"),
            ("FS16 文件不可见", "verify playback file 失败。", "不要尝试 start; 修文件名/路径/FS16 目录。"),
            ("FS16 load 失败", "SYST:ERR? 显示 corrupt/missing 或设备错误。", "保留错误、手动在 FS16 UI 打开 .smu。"),
            ("FS16 start 失败", "load 成功但 start_emulation 失败。", "记录 DIAG:SIMU:STATe? 和 SYST:ERR?, 检查 start 命令模板。"),
            ("KPI 无数据", "FS16 已正常但 mock attach/KPI 失败。", "先看 mock BS 状态和 sequence 参数; 不把它解读为真实射频问题。"),
        ],
        [1900, 3550, 3910],
        font_size=8.6,
    )

    add_heading(doc, "6. 现场记录表", 1)
    add_table(
        doc,
        ["项目", "现场记录"],
        [
            ("FS16 endpoint", ""),
            ("FS16 *IDN? 返回", ""),
            ("MMEM:CDIR? 返回", ""),
            ("目标 .smu 文件名/完整路径", ""),
            ("fs16_playback_smoke load-only 结果", ""),
            ("fs16_playback_smoke start/stop 结果", ""),
            ("fs16_hybrid_kpi_smoke 结果", ""),
            ("KPI summary 截图/诊断 run id", ""),
            ("异常 SYST:ERR? / DIAG:SIMU:STATe?", ""),
            ("下一步需要改的软件点", ""),
        ],
        [3000, 6360],
        font_size=9.2,
    )

    add_heading(doc, "7. 关键源码索引", 1)
    add_table(
        doc,
        ["用途", "文件"],
        [
            ("HAL real/mock 决策", "api-service/app/services/instrument_hal_service.py"),
            ("driver mode API", "api-service/app/api/instrument.py"),
            ("GUI driver mode 与 FS16 文件参数", "gui/src/App.tsx"),
            ("FS16 playback driver", "api-service/app/hal/propsim_fs16_playback.py"),
            ("Hybrid KPI sequence", "api-service/app/diagnostics/sequences/fs16_hybrid_kpi_smoke.py"),
            ("Sequence Runner 结果卡", "gui/src/features/Diagnostics/SequenceRunnerPanel.tsx"),
        ],
        [2800, 6560],
        font_size=9.0,
    )

    doc.add_page_break()
    add_heading(doc, "附录: 出发前软件测试记录", 1)
    add_table(
        doc,
        ["命令", "结果"],
        [
            (".venv/bin/python -m pytest tests/test_hal_mode_force_mock.py tests/test_fs16_hybrid_kpi_sequence.py tests/test_propsim_fs16_playback_driver.py -q", "32 passed, 61 warnings。"),
            ("npm run build", "构建成功; 仅有既有 Vite dynamic import / chunk size warning。"),
        ],
        [5200, 4160],
        font_size=8.5,
    )
    add_callout(
        doc,
        "本文件的边界",
        "这份计划确认的是软件控制流程和现场操作步骤。真实射频性能结论必须等真实基站、真实 DUT、校准链路和正式测试计划接入后再判断。",
        fill=CAUTION,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build_doc()
    print(path)
